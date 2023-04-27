import io
import os
import re
import tempfile
from datetime import date
from io import BytesIO

import boto3
from docx import Document
import openai
import streamlit as st
import tiktoken
from PIL import Image
from google.oauth2 import service_account
from googleapiclient.discovery import build
from pdf2image import convert_from_path
from PyPDF2 import PdfReader

# Langchain libraries
from langchain.chains import LLMChain
from langchain.document_loaders import TextLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma

# Initialize session state attributes
if "letter_text" not in st.session_state:
    st.session_state["letter_text"] = ""

if "faq_text" not in st.session_state:
    st.session_state["faq_text"] = ""

if "output" not in st.session_state:
    st.session_state["output"] = ""

if "doc_ID" not in st.session_state:
    st.session_state["doc_ID"] = ""

if "generated_output" not in st.session_state:
    st.session_state["generated_output"] = ""

# API keys and creds
openai.api_key=st.secrets["OPENAI_API_KEY"]
client = boto3.client('textract', region_name=st.secrets["AWS_DEFAULT_REGION"], aws_access_key_id=st.secrets["AWS_ACCESS_KEY_ID"], aws_secret_access_key=st.secrets["AWS_SECRET_ACCESS_KEY"])

credentials = service_account.Credentials.from_service_account_info(
    st.secrets["gcp_service_account"],
    scopes=[
        "https://www.googleapis.com/auth/drive",
        "https://www.googleapis.com/auth/documents",
    ],
)

drive_service = build('drive', 'v3', credentials=credentials)
docs_service = build('docs', 'v1', credentials=credentials)

# page title
st.set_page_config(layout="wide", page_icon="🏛", page_title="ParliamentGPT")
st.title("🖋️ AutoScribe")
st.markdown(
    """
    AutoScribe reads letters and drafts responses based on your FAQs, all in one click! Simply upload your FAQs and the correspondence you are responding to and click generate. You can even download the output in your own template by creating a Google Docs template and sharing it with parliamentgpt@gmail.com. Just don't forget to proofread and edit (AutoScribe is in beta so sometimes gets things wrong!)
    """
)

# general functions
def image_to_text(image):
    with BytesIO() as output:
        image.save(output, format="JPEG")
        img_data = output.getvalue()
    response = client.detect_document_text(
        Document={'Bytes': img_data}
    )
    text = ""
    for item in response["Blocks"]:
        if item["BlockType"] == "LINE":
            print(item["Text"])
            text = text + " " + item["Text"]
    return text

def convert_to_text(file):
    # Determine the file type
    file_type = os.path.splitext(file.name)[-1].lower()

    supported_image_formats = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif', '.webp']
    supported_text_formats = ['.txt', '.eml', '.html', '.rtf', '.md', '.csv', '.json', '.xml', '.doc', '.docx']

    # Convert PDF to text
    if file_type == '.pdf':
        file_reader = PdfReader(file)
        file_text = ""
        for page_num in range(len(file_reader.pages)):
            page = file_reader.pages[page_num]
            file_text += page.extract_text()
        if not file_text.strip():
            images = convert_from_path(pdf_file)
            for image in images:
                file_text += image_to_text(image)
        return file_text

    if file_type in supported_text_formats:
        with open(file, 'rb') as file:
            if file_type == '.docx':
                text = docx2txt.process(file)
            elif file_type == '.doc':
                doc = Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                text = file.read()
        return text

    # Convert image to text
    elif file_type in supported_image_formats:
        with Image.open(file, 'r') as file:
            text = image_to_text(file)
        return text
    
    # Unsupported file type
    else:
        raise ValueError("Unsupported file type.")

def string_to_temp_txt_file(content):
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as temp_file:
        temp_file.write(content)
        return temp_file.name

#@st.cache_data(ttl=None, max_entries=None, show_spinner=True, persist=None)
def generate_index(faq_text):
    temp_file_path = string_to_temp_txt_file(faq_text)
    loader = TextLoader(temp_file_path)
    documents = loader.load()
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=0)
    texts = text_splitter.split_documents(documents)
    embeddings = OpenAIEmbeddings()
    index = Chroma.from_documents(texts, embeddings)
    return index

#@st.cache_data(ttl=None, max_entries=None, show_spinner=True, persist=None)
def generate_output(index, query, word_count):
    prompt_template = """You are a UK politician responding to a letter from a constituent. First, convert the constituent's letter to plaintext then use context from your FAQs to draft a response. Address the constituent by their name, leave a space between each paragraph, and do not include your name at the end. The word count is {word_count} words.
    Context: {context}
    Letter: {letter}
    Answer:"""

    PROMPT = PromptTemplate(
        template=prompt_template, input_variables=["context", "letter", "word_count"]
    )
    llm = OpenAI(temperature=0)
    chain = LLMChain(llm=llm, prompt=PROMPT)

    docs = index.similarity_search(query, k=1)
    inputs = [{"context": doc.page_content, "letter": query, "word_count": str(word_count)} for doc in docs]
    output = chain.apply(inputs)[0]['text']
    return output

# streamlit button functions
def print_output():
    article = st.session_state.generated_output
    return article

#@st.cache_data(ttl=None, max_entries=None, show_spinner=True, persist=None)
def generate_address(letter):
    response = openai.ChatCompletion.create(
        model = "gpt-3.5-turbo",
        messages=[
            #{"role": "system", "content": "You receive text that has been extracted from an image using pytesseract, however the OCR is not perfect."},
            {"role": "user", "content":"Take the following text, extract the address, and write it out, with each part on a new line. If you can't find an address leave it blank:" + letter},
        ]
    )
    constituent_address = response['choices'][0]['message']['content']
    return constituent_address

def generate_doc(doc_ID, letter, output):
    c_address = generate_address(letter)

    # Use the global drive_service and docs_service objects
    global drive_service, docs_service

    # Use the Google Drive API to create a copy of the document
    copy_request = drive_service.files().copy(fileId=doc_ID)
    copied_doc = copy_request.execute()
    document_id = copied_doc['id']

    # Use the Google Docs API to retrieve the copied document
    doc = docs_service.documents().get(documentId=document_id).execute()

    # Define the text to replace and its replacement
    replace_dict = {
        '{{constituent_address}}': c_address,
        '{{letter_content}}': output,
        '{{date}}': date.today().strftime('%-d %B %Y'),
    }

    # Make edits to the copied document (e.g. replace some text)
    requests = []
    for find, replace in replace_dict.items():
        requests.append({
            'replaceAllText': {
                'containsText': {
                    'text': find,
                    'matchCase': True
                },
                'replaceText': replace
            }
        })

    result = docs_service.documents().batchUpdate(documentId=document_id, body={'requests': requests}).execute()

    # Use the Google Drive API to export the edited copied document as a Microsoft Word document
    file_content = drive_service.files().export(fileId=document_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document').execute()

    # Delete the copied file after downloading it
    drive_service.files().delete(fileId=document_id).execute()

    # Save the Word document content in memory
    doc = Document(io.BytesIO(file_content))
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes

def extract_google_docs_id(doc_ID: str) -> str:
    pattern = r'd/([\w-]+)/'
    match = re.search(pattern, doc_ID)
    if match:
        return match.group(1)
    else:
        st.warning('Add template URL and press enter', icon="⚠️")
        # raise ValueError("Add Google Docs URL")

encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")

def count_tokens(string: str) -> int:
    tokens = len(encoding.encode(string))
    return tokens

# layout

tab1, tab2 = st.tabs(["AutoScribe", "Instructions"])

with tab1:

    col1, col2 = st.columns([1, 1])

    with col1:

        faq = st.file_uploader("FAQs")
        if faq is not None:
            with st.spinner('Processing'):
                if st.session_state["faq_text"] != convert_to_text(faq):
                    st.session_state["faq_text"] = convert_to_text(faq)
                    st.session_state["index"] = generate_index(st.session_state.faq_text)
            st.success('FAQs uploaded!', icon="✅")

        letter = st.file_uploader("Upload letter")
        if letter is not None:
            with st.spinner('Processing'):
                if st.session_state["letter_text"] != convert_to_text(letter):
                    st.session_state["letter_text"] = convert_to_text(letter)
            st.success('Letter uploaded!', icon="✅")

            faq_token_count = count_tokens(st.session_state.faq_text)
            letter_token_count = count_tokens(st.session_state.letter_text)
            output_token_count = count_tokens(st.session_state.generated_output)

            if faq_token_count + letter_token_count + output_token_count >= 3800:
                st.warning('Your FAQ document or letter may be too long for the demo. A total of three pages between them is ideal', icon="⚠️")

        word_count = st.slider("Choose a word count:", min_value = 200, max_value = 400, value = 250)

        doc_ID = st.text_input(label="Google Docs ID:", value = "https://docs.google.com/document/d/1qhd6sbq5aTeBz3WGATrgyEZ1X456mWsz3ZsJywngx4Q/edit") 
        if doc_ID is not None:
            if st.session_state["doc_ID"] != extract_google_docs_id(doc_ID):
                st.session_state["doc_ID"] = extract_google_docs_id(doc_ID)

    with col2:
      
      output = st.text_area(label="Generated output:", height=680, key="output", value=st.session_state.generated_output)
        if output is not None:
            if st.session_state["generated_output"] != output:
                st.session_state["generated_output"] = output

        g, c, d = st.columns([1, 1, 1])

        with g:
            generate_button = st.button("Generate")
            if generate_button:
                if hasattr(st.session_state, "faq_text") and hasattr(st.session_state, "letter_text"):
                    with st.spinner('Generating...'):
                        if st.session_state["generated_output"] != generate_output(st.session_state["index"], st.session_state["letter_text"], word_count):
                            st.session_state["generated_output"] = generate_output(st.session_state["index"], st.session_state["letter_text"], word_count)

        with c:
            clear_button = st.button("Clear")
            if clear_button:
                st.session_state.generated_output = ""

        with d:
            if st.session_state.generated_output:
                st.download_button(
                    label="Download",
                    data=generate_doc(st.session_state.doc_ID, st.session_state.letter_text, st.session_state.generated_output).getvalue(),
                    file_name="draft_letter.docx",
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

with tab2:
    st.markdown(
        """
        1. ℹ️ **Upload your FAQs.** Here is a [sample](https://docs.google.com/document/d/1qZ3mjCUtM_DudBfPx1hZg6c8-G_5yF0YmTOBRerpcag/edit) of what this could look like. You don't need to worry about copying and pasting specific FAQs depending on the topic, AutoScribewill find the right FAQ from the whole document.
        2. 📬 **Upload the letter you are responding to.** Here is a [sample](https://docs.google.com/document/d/1P9bmBsapYkklPTa3XX8wRuYeeuUgnGq8XUvL_IX3PI0/edit) of what this could look like. Autoscribe can extract text from images or scanned, even if it's handwritten!
        3. 🎛 **Adjust the word count.** You can tweak the settings using the slider.
        4. 🤖 **Click generate!**  Your letter will be displayed in the output box on the right.
        5. 🤓 **Proofread and edit**. Don't forget that AutoScribe is just a tool, and sometimes it gets things wrong! Don't forget to proofread and edit your output before using it.
        
        [Optional] If you would like to download the output in your own template, you can follow the additional steps below.
        1. 📝 **Create a letter template in Google Docs.** The template must include placeholders {{letter_content}}, {{respondent_address}}, {{date}} where you want the content, respondent's address and date to go. For an example, see [here](https://docs.google.com/document/d/17la5aNiLcFGdk43JrTvXBVVW9561Xuc0s0896pczUlU/edit).
        2. 🔖 **Share the Google Doc with [ParliamentGPT@gmail.com](parliamentgpt@gmail.com).** Click Share in the top right of the Google Doc to share. Also change access to that "anyone with the link can view."
        3. 📎 **Paste a link to your template in the box labeled "template".** Click Share in the top right of the Google Doc to get a link, or just copy and paste the url from your browser. 
        
        If you are having any issues, read through the instructions again carefully. If you still can't solve your problem, email [parliamentgpt@gmail.com](parliamentgpt@gmail.com) and we'll get back to you.
        """
    )
